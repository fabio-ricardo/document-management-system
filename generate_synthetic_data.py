import os
from faker import Faker
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Initialize Faker for generating fake data
fake = Faker()

# Define the folder structure
data_folder = "synthetic_data"
os.makedirs(data_folder, exist_ok=True)

# Function to generate a synthetic invoice (PDF)
def generate_invoice_pdf(file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    c.setFont("Helvetica", 12)

    # Add invoice header
    c.drawString(100, 750, "Invoice")
    c.drawString(100, 730, f"Invoice Number: {fake.random_int(min=1000, max=9999)}")
    c.drawString(100, 710, f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    c.drawString(100, 690, f"Bill To: {fake.company()}")

    # Add invoice items
    c.drawString(100, 650, "Item Description          Quantity          Price")
    y = 630
    for _ in range(3):
        c.drawString(100, y, f"{fake.word()}          {fake.random_int(min=1, max=10)}          ${fake.random_int(min=10, max=100)}")
        y -= 20

    # Add total
    c.drawString(100, y - 20, f"Total: ${fake.random_int(min=50, max=500)}")

    c.save()

# Function to generate a synthetic contract (DOCX)
def generate_contract_docx(file_path):
    doc = Document()
    doc.add_heading("Contract Agreement", 0)

    # Add contract details
    doc.add_paragraph(f"This agreement is made and entered into on {datetime.now().strftime('%Y-%m-%d')} by and between:")
    doc.add_paragraph(f"Party A: {fake.company()}")
    doc.add_paragraph(f"Party B: {fake.company()}")
    doc.add_paragraph("Terms and Conditions:")
    for _ in range(5):
        doc.add_paragraph(fake.sentence())

    doc.add_paragraph("Signatures:")
    doc.add_paragraph(f"Party A: ___________________________")
    doc.add_paragraph(f"Party B: ___________________________")

    doc.save(file_path)

# Function to generate a synthetic report (PDF)
def generate_report_pdf(file_path):
    c = canvas.Canvas(file_path, pagesize=letter)
    c.setFont("Helvetica", 12)

    # Add report header
    c.drawString(100, 750, "Monthly Report")
    c.drawString(100, 730, f"Report Date: {datetime.now().strftime('%Y-%m-%d')}")
    c.drawString(100, 710, f"Prepared By: {fake.name()}")

    # Add report content
    c.drawString(100, 680, "Summary:")
    y = 660
    for _ in range(5):
        c.drawString(100, y, fake.sentence())
        y -= 20

    c.save()

# Generate synthetic documents
generate_invoice_pdf(os.path.join(data_folder, "invoice.pdf"))
generate_contract_docx(os.path.join(data_folder, "contract.docx"))
generate_report_pdf(os.path.join(data_folder, "report.pdf"))

# Create a README file for the synthetic data
readme = """
# Synthetic Data for Testing

This folder contains synthetic documents generated for testing the Document Management System.

## Files

1. **invoice.pdf**: A synthetic invoice document.
2. **contract.docx**: A synthetic contract document.
3. **report.pdf**: A synthetic report document.

## Usage

These documents can be used to test the document upload, categorization, and summarization features of the application.
"""
with open(os.path.join(data_folder, "README.md"), "w") as f:
    f.write(readme)

print("Synthetic data generated successfully!")